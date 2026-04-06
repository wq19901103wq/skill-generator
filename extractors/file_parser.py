#!/usr/bin/env python3
"""
文件解析器 - 支持多种文件格式
"""

import os
import json
import re
from pathlib import Path
from typing import Dict, List, Optional, Union


class FileParser:
    """通用文件解析器"""
    
    SUPPORTED_FORMATS = {
        '.txt': 'text',
        '.md': 'text',
        '.json': 'json',
        '.csv': 'csv',
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'docx',
    }
    
    def __init__(self):
        self.text_content = ""
        self.metadata = {}
        
    def parse(self, file_path: Union[str, Path]) -> Dict:
        """
        解析文件，返回内容和元数据
        
        Returns:
            {
                'content': str,  # 文本内容
                'metadata': dict,  # 文件元数据
                'format': str,  # 文件格式
            }
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        ext = file_path.suffix.lower()
        format_type = self.SUPPORTED_FORMATS.get(ext, 'unknown')
        
        self.metadata = {
            'filename': file_path.name,
            'format': format_type,
            'size': file_path.stat().st_size,
        }
        
        # 根据格式选择解析方法
        if format_type == 'text':
            content = self._parse_text(file_path)
        elif format_type == 'json':
            content = self._parse_json(file_path)
        elif format_type == 'csv':
            content = self._parse_csv(file_path)
        elif format_type == 'pdf':
            content = self._parse_pdf(file_path)
        elif format_type == 'docx':
            content = self._parse_docx(file_path)
        else:
            # 尝试作为文本读取
            try:
                content = self._parse_text(file_path)
            except:
                content = f"[无法解析文件: {file_path.name}]"
        
        self.text_content = content
        
        return {
            'content': content,
            'metadata': self.metadata,
            'format': format_type,
        }
    
    def parse_multiple(self, file_paths: List[Union[str, Path]]) -> Dict[str, Dict]:
        """解析多个文件"""
        results = {}
        for path in file_paths:
            try:
                results[str(path)] = self.parse(path)
            except Exception as e:
                results[str(path)] = {'error': str(e)}
        return results
    
    def _parse_text(self, file_path: Path) -> str:
        """解析文本文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # 如果都失败，使用 latin-1 读取（不会抛异常）
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()
    
    def _parse_json(self, file_path: Path) -> str:
        """解析 JSON 文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # 将 JSON 转换为可读的文本格式
        if isinstance(data, dict):
            return self._dict_to_text(data)
        elif isinstance(data, list):
            return '\n\n'.join(self._dict_to_text(item) if isinstance(item, dict) else str(item) 
                             for item in data)
        else:
            return str(data)
    
    def _parse_csv(self, file_path: Path) -> str:
        """解析 CSV 文件"""
        import csv
        
        content = []
        encodings = ['utf-8', 'gbk', 'gb2312']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, newline='') as f:
                    reader = csv.DictReader(f)
                    rows = list(reader)
                    
                    # 转换为文本
                    if rows:
                        for i, row in enumerate(rows, 1):
                            content.append(f"记录 {i}:")
                            for key, value in row.items():
                                if value:
                                    content.append(f"  {key}: {value}")
                            content.append("")
                
                return '\n'.join(content)
            except UnicodeDecodeError:
                continue
        
        return "[CSV 解析失败]"
    
    def _parse_pdf(self, file_path: Path) -> str:
        """解析 PDF 文件"""
        try:
            import PyPDF2
            
            text = []
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                self.metadata['pages'] = len(reader.pages)
                
                for page in reader.pages:
                    text.append(page.extract_text() or "")
            
            return '\n'.join(text)
        except ImportError:
            return "[需要安装 PyPDF2: pip install PyPDF2]"
        except Exception as e:
            return f"[PDF 解析错误: {e}]"
    
    def _parse_docx(self, file_path: Path) -> str:
        """解析 Word 文档"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_text:
                        paragraphs.append(' | '.join(row_text))
            
            return '\n'.join(paragraphs)
        except ImportError:
            return "[需要安装 python-docx: pip install python-docx]"
        except Exception as e:
            return f"[Word 解析错误: {e}]"
    
    def _dict_to_text(self, data: Dict, indent: int = 0) -> str:
        """将字典转换为文本"""
        lines = []
        prefix = "  " * indent
        
        for key, value in data.items():
            if isinstance(value, dict):
                lines.append(f"{prefix}{key}:")
                lines.append(self._dict_to_text(value, indent + 1))
            elif isinstance(value, list):
                lines.append(f"{prefix}{key}:")
                for item in value:
                    if isinstance(item, dict):
                        lines.append(self._dict_to_text(item, indent + 1))
                    else:
                        lines.append(f"{prefix}  - {item}")
            else:
                lines.append(f"{prefix}{key}: {value}")
        
        return '\n'.join(lines)
    
    @classmethod
    def detect_file_type(cls, file_path: Union[str, Path]) -> str:
        """检测文件类型"""
        ext = Path(file_path).suffix.lower()
        return cls.SUPPORTED_FORMATS.get(ext, 'unknown')
    
    @classmethod
    def is_supported(cls, file_path: Union[str, Path]) -> bool:
        """检查是否支持该文件格式"""
        return cls.detect_file_type(file_path) != 'unknown'
